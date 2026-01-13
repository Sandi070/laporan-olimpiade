import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import datetime
from google import genai 

# --- FUNGSI FORMAT TANGGAL INDONESIA ---
def tanggal_indo(tgl):
    hari_dict = {'Monday': 'Senin', 'Tuesday': 'Selasa', 'Wednesday': 'Rabu', 'Thursday': 'Kamis', 'Friday': 'Jumat', 'Saturday': 'Sabtu', 'Sunday': 'Minggu'}
    bulan_dict = {'January': 'Januari', 'February': 'Februari', 'March': 'Maret', 'April': 'April', 'May': 'Mei', 'June': 'Juni', 'July': 'Juli', 'August': 'Agustus', 'September': 'September', 'October': 'Oktober', 'November': 'November', 'December': 'Desember'}
    return f"{hari_dict[tgl.strftime('%A')]}, {tgl.strftime('%d')} {bulan_dict[tgl.strftime('%B')]} {tgl.strftime('%Y')}"

# --- KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Sistem Laporan MAN IC Kendari", layout="wide")

# --- JUDUL APLIKASI ---
st.title("🏫 Generator Laporan Kegiatan MAN IC Kendari")
st.markdown("Aplikasi Laporan Program Unggulan (Olimpiade, TKA, UTBK, Klinik, Ekskul, Pengasuhan, dll).")

# --- SIDEBAR: MENU UTAMA & API ---
with st.sidebar:
    st.header("⚙️ Pengaturan")
    api_key = st.text_input("Google Gemini API Key", type="password")
    if not api_key:
        st.warning("Masukkan API Key agar fitur AI aktif.")
    
    st.divider()
    
    st.header("📂 Pilih Program")
    jenis_program = st.selectbox(
        "Jenis Kegiatan Laporan:",
        [
            "Bimbingan Olimpiade",
            "Bimbingan TKA (Kompetensi Akademik)",
            "Bimbingan UTBK/SNBT",
            "Klinik Mata Pelajaran (Remedial)",
            "Ekstrakurikuler",
            "Karya Ilmiah Remaja (KIR)",
            "Pendampingan Belajar Malam",
            "Kegiatan Pengasuhan (Guru Asuh)"
        ]
    )

# --- INPUT DATA UTAMA (DINAMIS SESUAI PROGRAM) ---
st.header(f"📝 Input Data: {jenis_program}")

col1, col2 = st.columns(2)

# Variabel default
mapel = "-"
materi = "-"
kategori_malam = "-"
topik_pengasuhan = "-"

with col1:
    nama_pembahas = st.text_input("Nama Guru/Pembina", "Sandi Saputra, S.Pd.")
    
    # --- LOGIKA MATA PELAJARAN SPESIFIK ---
    
    # 1. OLIMPIADE (Tetap dengan 9 Mapel OSN)
    if jenis_program == "Bimbingan Olimpiade":
        list_mapel = ["Kimia", "Fisika", "Biologi", "Matematika", "Ekonomi", "Geografi", "Kebumian", "Astronomi", "Informatika"]
        mapel = st.selectbox("Bidang Olimpiade", list_mapel)
        materi = st.text_input("Materi Bahasan", "Contoh: Stoikiometri / Konsep Mol")

    # 2. TKA (Plus Sejarah, Sosiologi | Minus Kebumian, Astro, Info, TPS)
    elif jenis_program == "Bimbingan TKA (Kompetensi Akademik)":
        list_mapel = [
            "Matematika", "Fisika", "Kimia", "Biologi", 
            "Ekonomi", "Geografi", "Sejarah", "Sosiologi",
            "Bahasa Indonesia", "Bahasa Inggris"
        ]
        mapel = st.selectbox("Mata Pelajaran TKA", list_mapel)
        materi = st.text_input("Materi / Kompetensi", "Contoh: Latihan Soal TKA Paket 1")

    # 3. UTBK (Plus Penalaran Umum | Minus Kebumian, Astro, Info)
    elif jenis_program == "Bimbingan UTBK/SNBT":
        list_mapel = [
            "Penalaran Umum", "Pengetahuan Kuantitatif", "Pemahaman Bacaan & Menulis",
            "Pengetahuan & Pemahaman Umum", "Literasi Bahasa Indonesia", "Literasi Bahasa Inggris",
            "Penalaran Matematika", "Matematika", "Fisika", "Kimia", "Biologi", 
            "Ekonomi", "Geografi", "Sejarah", "Sosiologi"
        ]
        mapel = st.selectbox("Subtes / Mapel UTBK", list_mapel)
        materi = st.text_input("Topik Bahasan", "Contoh: Trik Cepat Pengerjaan Soal Penalaran")

    # 4. KLINIK (Paling Lengkap: Plus Mapel Agama)
    elif jenis_program == "Klinik Mata Pelajaran (Remedial)":
        list_mapel = [
            "Matematika", "Fisika", "Kimia", "Biologi", 
            "Ekonomi", "Geografi", "Sejarah", "Sosiologi",
            "Bahasa Indonesia", "Bahasa Inggris",
            "Akidah Akhlak", "Fiqih", "Bahasa Arab", 
            "Alquran Hadits", "Sejarah Kebudayaan Islam (SKI)"
        ]
        mapel = st.selectbox("Mata Pelajaran Klinik", list_mapel)
        materi = st.text_input("Materi Remedial", "Contoh: Pemahaman ulang Bab Termokimia")

    # 5. EKSTRAKURIKULER (Baru)
    elif jenis_program == "Ekstrakurikuler":
        list_ekskul = ["Robotik", "Pramuka", "PMR", "Jurnalistik"]
        mapel = st.selectbox("Bidang Ekstrakurikuler", list_ekskul) # Kita simpan nama ekskul di variabel mapel
        materi = st.text_input("Agenda Kegiatan", "Contoh: Latihan Baris Berbaris / Coding Arduino")

    # 6. KIR
    elif jenis_program == "Karya Ilmiah Remaja (KIR)":
        mapel = st.selectbox("Bidang Penelitian", ["IPA (Saintek)", "IPS (Soshum)", "Keagamaan", "Teknologi"])
        materi = st.text_input("Judul Penelitian Siswa", "Contoh: Pemanfaatan Limbah Sagu")

    # 7. BELAJAR MALAM
    elif jenis_program == "Pendampingan Belajar Malam":
        kategori_malam = st.selectbox("Jenis Kegiatan", ["Belajar Mandiri (KBM)", "Latihan Upacara", "Latihan Seni", "Lainnya"])
        materi = st.text_input("Detail Kegiatan", "Contoh: Persiapan Petugas Upacara")
        mapel = "Pendampingan Asrama"

    # 8. PENGASUHAN
    elif jenis_program == "Kegiatan Pengasuhan (Guru Asuh)":
        mapel = "Pengasuhan & Konseling"
        topik_pengasuhan = st.text_area("Topik Diskusi", "Contoh: Evaluasi kebersihan kamar dan motivasi belajar.")

with col2:
    tanggal = st.date_input("Tanggal Kegiatan", datetime.date.today())
    waktu_mulai = st.time_input("Waktu Mulai", datetime.time(16, 0)) # Default jam sore untuk ekskul
    waktu_selesai = st.time_input("Waktu Selesai", datetime.time(17, 30))
    jumlah_peserta = st.number_input("Jumlah Siswa Hadir", min_value=1, value=10)

# --- UPLOAD FOTO ---
st.subheader("📷 Dokumentasi")
uploaded_file = st.file_uploader("Upload Foto Kegiatan", type=['png', 'jpg', 'jpeg'])
if uploaded_file:
    st.image(uploaded_file, width=400, caption="Preview Foto")

# --- LOGIKA AI (PROMPT DINAMIS) ---
def generate_description_ai(api_key, program, mapel, materi, topik, kategori_malam, tanggal_obj, peserta, mulai, selesai):
    try:
        client = genai.Client(api_key=api_key)
        tgl_teks = tanggal_indo(tanggal_obj)
        
        # PROMPT ENGINEER
        konteks = ""
        
        if "Olimpiade" in program:
            konteks = f"Fokus: Bimbingan olimpiade {mapel} materi {materi}. Siswa dipersiapkan untuk kompetisi tingkat lanjut."
        elif "TKA" in program:
            konteks = f"Fokus: Persiapan Tes Kompetensi Akademik (TKA) mapel {mapel}. Drill soal dan penguatan konsep dasar."
        elif "UTBK" in program:
            konteks = f"Fokus: Persiapan UTBK/SNBT subtes {mapel}. Bahas trik mengerjakan soal '{materi}' dengan cepat dan tepat."
        elif "Klinik" in program:
            konteks = f"Fokus: Remedial/Klinik mapel {mapel}. Pendampingan khusus bagi siswa yang belum tuntas materi {materi}."
        elif "Ekstrakurikuler" in program:
            konteks = f"Fokus: Kegiatan pengembangan diri bidang {mapel} (Ekskul). Kegiatan: {materi}. Tekankan pada pengembangan skill, kerjasama tim, dan karakter siswa."
        elif "KIR" in program:
            konteks = f"Fokus: Bimbingan riset KIR bidang {mapel}. Topik: {materi}."
        elif "Malam" in program:
            konteks = f"Fokus: Pendampingan malam ({kategori_malam}). Kegiatan: {materi}. Memastikan ketertiban asrama."
        elif "Pengasuhan" in program:
            konteks = f"Fokus: Pembinaan guru asuh. Topik: {topik}. Memberi motivasi dan solusi masalah siswa."

        prompt = f"""
        Buatkan 1 paragraf laporan kegiatan formal untuk MAN Insan Cendekia.
        Program: {program}
        Waktu: {tgl_teks}, {mulai}-{selesai}.
        Detail: {konteks}
        
        Output: Langsung paragraf isi laporan. Bahasa baku, profesional, tanpa judul.
        """

        response = client.models.generate_content(
            model="gemini-flash-latest", 
            contents=prompt
        )
        return response.text
        
    except Exception as e:
        return f"Gagal Generate. Error: {str(e)}"

# --- TOMBOL GENERATE ---
st.divider()
col_btn, col_res = st.columns([1, 2])

with col_btn:
    st.subheader("Langkah 1: Generate Teks")
    if st.button("✨ Buat Laporan Otomatis"):
        if api_key:
            with st.spinner("AI sedang berpikir..."):
                res = generate_description_ai(api_key, jenis_program, mapel, materi, topik_pengasuhan, kategori_malam, tanggal, jumlah_peserta, waktu_mulai, waktu_selesai)
                st.session_state['deskripsi'] = res
        else:
            st.error("API Key belum diisi.")

with col_res:
    deskripsi_final = st.text_area("Hasil Laporan:", value=st.session_state.get('deskripsi', ''), height=200)

# --- FUNGSI WORD ---
def create_docx(nama, program, tgl_obj, label_baris3, isi_baris3, deskripsi, gambar):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Header
    tahun = tgl_obj.year
    judul_prog = program.upper().split('(')[0].strip() # Bersihkan nama program
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"LAPORAN KEGIATAN\nPROGRAM {judul_prog}\nMAN INSAN CENDEKIA KOTA KENDARI TAHUN {tahun}")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Metadata
    tgl_indo = tanggal_indo(tgl_obj).upper()
    table = doc.add_table(rows=3, cols=3)
    
    metadata = [
        ("NAMA PEMBINA", ":", nama),
        ("HARI/TANGGAL", ":", tgl_indo),
        (label_baris3, ":", isi_baris3.upper())
    ]
    
    for i, data in enumerate(metadata):
        r = table.rows[i]
        r.cells[0].text = data[0]
        r.cells[1].text = data[1]
        r.cells[2].text = data[2]
        r.cells[0].paragraphs[0].runs[0].bold = True 
        
    doc.add_paragraph() 

    # Isi
    doc.add_paragraph("DESKRIPSI KEGIATAN").runs[0].bold = True
    doc.add_paragraph(deskripsi)
    doc.add_paragraph() 
    
    # Dokumentasi
    doc.add_paragraph("DOKUMENTASI").runs[0].bold = True
    if gambar:
        doc.add_picture(gambar, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- TOMBOL DOWNLOAD ---
st.subheader("Langkah 2: Download File")

if st.button("💾 Download Dokumen (.docx)"):
    if deskripsi_final:
        # Menentukan Label untuk baris ke-3 tabel Metadata
        label_meta = "MATERI/KEGIATAN"
        isi_meta = f"{mapel}: {materi}"
        
        if "Pengasuhan" in jenis_program:
            label_meta = "TOPIK PENGASUHAN"
            isi_meta = "PENGASUHAN SISWA"
        elif "Malam" in jenis_program:
            label_meta = "KEGIATAN"
            isi_meta = f"{kategori_malam} ({materi})"
        elif "KIR" in jenis_program:
            label_meta = "JUDUL PENELITIAN"
            isi_meta = f"KIR {mapel}: {materi}"
        elif "Ekstrakurikuler" in jenis_program:
            label_meta = "BIDANG EKSKUL"
            isi_meta = f"{mapel} ({materi})" # mapel berisi nama ekskul (Robotik dll)

        file_docx = create_docx(nama_pembahas, jenis_program, tanggal, label_meta, isi_meta, deskripsi_final, uploaded_file)
        
        file_name = f"Laporan_{jenis_program}_{tanggal}.docx"
        st.download_button("Klik untuk Unduh", file_docx, file_name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.success(f"File {jenis_program} berhasil dibuat!")
    else:
        st.error("Generate deskripsi dulu.")
